# -*- coding: utf-8 -*-
"""
পরীক্ষার কক্ষের রোল নাম্বার সিটিং শিট - ওয়েব অ্যাপ ভার্সন (Streamlit)
======================================================================
এটা আগের seating_sheet.py স্ক্রিপ্টের মতোই কাজ করে, শুধু পার্থক্য হলো
এখন রোল নাম্বার/কক্ষ নং ইত্যাদি কোড এডিট না করে সরাসরি ব্রাউজারে
একটা ফর্মে বসিয়ে দেওয়া যাবে, আর বাটনে ক্লিক করলেই docx ডাউনলোড হবে।

এই ভার্সনে নতুন যা যোগ হলো:
  ১. প্রতিটা কলাম-গ্রুপে সাব-কলাম সংখ্যা কাস্টম করা যাবে (অপশনাল, ডিফল্ট ২)
  ২. কিছু রোল নাম্বার "অনুপস্থিত" হিসেবে বাদ দেওয়া যাবে (অপশনাল) -
     বাদ দেওয়া রোলের জায়গায় কোনো বক্স থাকবে না, পরের রোল উপরে উঠে আসবে।

চালানোর নিয়ম (টার্মিনালে):
    streamlit run app.py
"""

import io
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import streamlit as st

_BOX_ID_COUNTER = [1000]


# ============================================================
# নিচের ফাংশনগুলো হুবহু আগের স্ক্রিপ্ট থেকে নেওয়া -- কোনো পরিবর্তন নাই
# ============================================================

def set_cell_border_none(cell):
    """টেবিলের সেল বর্ডার সম্পূর্ণ অদৃশ্য করে দেয় (গ্রিড লাইন দেখা যাবে না)।"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'nil')
        borders.append(tag)
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        shd.addprevious(borders)
    else:
        tcPr.append(borders)


def make_rounded_box(paragraph, text, width_cm=2.6, height_cm=1.0,
                      line_color="2E5BDA", fill_color="F4F8FF",
                      font_size=14, font_color="1A1A1A", bold=True):
    """একটা প্যারাগ্রাফের ভেতরে রাউন্ডেড রেক্টাঙ্গেল শেপ বসায় (VML)।"""
    _BOX_ID_COUNTER[0] += 1
    shape_id = _BOX_ID_COUNTER[0]
    width_pt = width_cm * 28.3465
    height_pt = height_cm * 28.3465

    xml = f'''
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:w10="urn:schemas-microsoft-com:office:word">
      <w:pict>
        <v:roundrect id="box{shape_id}" o:spid="_x0000_s{shape_id}"
            style="width:{width_pt}pt;height:{height_pt}pt"
            arcsize="25%" fillcolor="#{fill_color}" strokecolor="#{line_color}" strokeweight="1.5pt"
            xmlns:o="urn:schemas-microsoft-com:office:office">
          <v:textbox style="mso-fit-shape-to-text:t" inset="1pt,1pt,1pt,1pt">
            <w:txbxContent>
              <w:p>
                <w:pPr><w:jc w:val="center"/></w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:b w:val="{ "1" if bold else "0" }"/>
                    <w:color w:val="{font_color}"/>
                    <w:sz w:val="{font_size * 2}"/>
                  </w:rPr>
                  <w:t>{text}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:roundrect>
      </w:pict>
    </w:r>
    '''
    run_element = parse_xml(xml.strip())
    paragraph._p.append(run_element)


def add_header_bar(cell, text):
    """নীল রঙের হেডার বার বানায়।"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "Nirmala UI"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Nirmala UI')
    rFonts.set(qn('w:hAnsi'), 'Nirmala UI')
    rFonts.set(qn('w:cs'), 'Nirmala UI')
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '4472C4')
    tcPr.append(shd)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _fix_zoom_percent(docx_bytes: bytes) -> bytes:
    """settings.xml এর zoom percent স্কিমা-নিটপিক ঠিক করে দেয়।"""
    import zipfile, re
    src = io.BytesIO(docx_bytes)
    dst = io.BytesIO()
    with zipfile.ZipFile(src, 'r') as zin:
        with zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/settings.xml':
                    text = data.decode('utf-8')
                    if 'w:percent' not in text:
                        text = re.sub(
                            r'(<w:zoom\b)(?![^>]*w:percent)',
                            r'\1 w:percent="100"',
                            text
                        )
                    data = text.encode('utf-8')
                zout.writestr(item, data)
    return dst.getvalue()


def parse_absent_rolls(text: str) -> set:
    """
    "102175-102180,102185,102190-102200" এর মতো স্ট্রিং থেকে অনুপস্থিত
    রোল নাম্বারের সেট বানায়। কমা দিয়ে আলাদা আলাদা আইটেম, প্রতিটা আইটেম
    হয় একটা একক রোল (যেমন 102185), অথবা হাইফেন দেওয়া রেঞ্জ (যেমন 102175-102180,
    এক্ষেত্রে দুই প্রান্তসহ মাঝের সব রোল ধরা হবে)।
    """
    result = set()
    for part in text.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start_s, end_s = part.split("-", 1)
            start_n = int(start_s.strip())
            end_n = int(end_s.strip())
            if start_n > end_n:
                start_n, end_n = end_n, start_n
            result.update(range(start_n, end_n + 1))
        else:
            result.add(int(part))
    return result


def generate_present_rolls(start_roll: int, total_students: int, absent_rolls: set) -> tuple:
    """
    start_roll থেকে শুরু করে ধারাবাহিকভাবে রোল নাম্বার বসায়, কিন্তু absent_rolls
    সেটে থাকা নাম্বারগুলো বাদ দিয়ে এগিয়ে যায় (সেই জায়গায় কোনো গ্যাপ থাকে না,
    পরের উপস্থিত রোল উপরে উঠে আসে)। মোট total_students সংখ্যক উপস্থিত রোল
    পাওয়া পর্যন্ত চলতে থাকে।
    রিটার্ন করে: (উপস্থিত রোলের লিস্ট, সর্বশেষ রোল নাম্বার যেখানে থামল)
    """
    rolls = []
    current = start_roll
    while len(rolls) < total_students:
        if current not in absent_rolls:
            rolls.append(current)
        current += 1
    return rolls, current - 1


def build_document(room_no: str, start_roll: int, total_students: int,
                    group_counts: list, group_subcols: list,
                    absent_rolls: set, gap_width_cm: float) -> tuple:
    """
    মূল ফাংশন: ইনপুট নিয়ে একটা সম্পূর্ণ docx ফাইল মেমোরিতে বানিয়ে
    bytes আকারে রিটার্ন করে।

    group_counts  -> প্রতিটা কলাম-গ্রুপে কতজন উপস্থিত পরীক্ষার্থীর রোল বসবে
    group_subcols -> প্রতিটা কলাম-গ্রুপ কয়টা সাব-কলামে ভাগ হবে (যেমন 2,2,3)
    absent_rolls  -> যেসব রোল নাম্বার বাদ দিতে হবে (set)
    """
    _BOX_ID_COUNTER[0] = 1000  # প্রতিবার নতুন করে শুরু করা, আইডি কনফ্লিক্ট এড়াতে

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(f"কক্ষ {room_no}")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(f"মোট পরীক্ষার্থীঃ {total_students} জন")
    run2.font.size = Pt(11)

    # ধারাবাহিক রোল লিস্ট বানানো (absent বাদ দিয়ে), তারপর গ্রুপ অনুযায়ী ভাগ করা
    all_rolls, last_roll = generate_present_rolls(start_roll, total_students, absent_rolls)
    col_rolls = []
    cursor = 0
    for count in group_counts:
        col_rolls.append(all_rolls[cursor:cursor + count])
        cursor += count

    n_groups = len(group_counts)
    # প্রতিটা গ্রুপে কয় রো লাগবে (সাব-কলাম সংখ্যা অনুযায়ী)
    rows_per_group = []
    for rolls, subcols in zip(col_rolls, group_subcols):
        rows_per_group.append(-(-len(rolls) // subcols))  # ceil division
    max_rows = max(rows_per_group) if rows_per_group else 0

    n_gaps = n_groups - 1
    n_cols = sum(group_subcols) + n_gaps
    # +1 অতিরিক্ত রো যোগ করা হলো হেডারের ঠিক নিচে একটা ফাঁকা স্পেসার রো-এর জন্য
    table = doc.add_table(rows=2 + max_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    group_col_start = []
    gap_col_indices = []
    pos = 0
    for g in range(n_groups):
        group_col_start.append(pos)
        pos += group_subcols[g]
        if g < n_groups - 1:
            gap_col_indices.append(pos)
            pos += 1

    col_width = Cm(3.0)
    gap_width = Cm(gap_width_cm)
    for col_idx, col in enumerate(table.columns):
        width = gap_width if col_idx in gap_col_indices else col_width
        for cell in col.cells:
            cell.width = width

    header_row = table.rows[0]
    BENGALI_DIGITS = str.maketrans("0123456789", "০১২৩৪৫৬৭৮৯")
    for g in range(n_groups):
        start = group_col_start[g]
        subcols = group_subcols[g]
        merged = header_row.cells[start]
        for k in range(1, subcols):
            merged = merged.merge(header_row.cells[start + k])
        label_num = str(g + 1).translate(BENGALI_DIGITS)
        add_header_bar(merged, f"কলাম {label_num}")
        set_cell_border_none(merged)
    for gap_idx in gap_col_indices:
        gap_cell = header_row.cells[gap_idx]
        gap_cell.text = ""
        set_cell_border_none(gap_cell)

    # হেডারের ঠিক নিচের রো-টা ফাঁকা স্পেসার হিসেবে রাখা (ছোট উচ্চতা, কোনো বর্ডার/টেক্সট নাই)
    spacer_row = table.rows[1]
    spacer_row.height = Cm(0.4)
    for cell in spacer_row.cells:
        cell.text = ""
        set_cell_border_none(cell)

    for r in range(max_rows):
        row = table.rows[r + 2]
        for gap_idx in gap_col_indices:
            gap_cell = row.cells[gap_idx]
            gap_cell.text = ""
            set_cell_border_none(gap_cell)

        for g in range(n_groups):
            rolls = col_rolls[g]
            subcols = group_subcols[g]
            start = group_col_start[g]
            per_subcol = -(-len(rolls) // subcols)  # প্রতি সাব-কলামে কয়টা রোল (ceil)

            # column-major ক্রম: প্রথম সাব-কলাম আগে উপর-নিচ পূর্ণ হয়, তারপর পরের সাব-কলাম...
            for s in range(subcols):
                cell = row.cells[start + s]
                set_cell_border_none(cell)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                idx = s * per_subcol + r
                if idx < len(rolls):
                    make_rounded_box(p, str(rolls[idx]))

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = _fix_zoom_percent(buf.getvalue())
    return docx_bytes, last_roll


# ============================================================
# ওয়েব ফর্ম (UI) -- এখানেই আপনি ব্রাউজারে যা দেখবেন তার বর্ণনা
# ============================================================

st.set_page_config(page_title="সিটিং শিট জেনারেটর", page_icon="📋")

st.title("📋 পরীক্ষার সিটিং শিট জেনারেটর")
st.write("নিচে তথ্য দিয়ে বাটনে ক্লিক করলেই Word ফাইল তৈরি হয়ে যাবে।")

col1, col2 = st.columns(2)
with col1:
    room_no = st.text_input("কক্ষ নম্বর", value="১১০")
    start_roll = st.number_input("শুরুর রোল নাম্বার", min_value=1, value=102169, step=1)
with col2:
    total_students = st.number_input("মোট পরীক্ষার্থী (উপস্থিত)", min_value=1, value=40, step=1)
    gap_width = st.number_input("কলামের মাঝে ফাঁকা জায়গা (cm)", min_value=0.0, value=1.0, step=0.5)

st.subheader("কলাম বিন্যাস")
st.caption(
    "প্রতিটা কলাম-গ্রুপে কতজন (উপস্থিত) পরীক্ষার্থী থাকবে লিখুন, কমা দিয়ে আলাদা করে। "
    "যেমন ৪০ জনের জন্য ৩টা গ্রুপ: 12,12,16"
)
group_counts_text = st.text_input("কলাম-প্রতি পরীক্ষার্থী সংখ্যা", value="12,12,16")

st.caption(
    "🔧 অপশনাল: প্রতিটা কলাম-গ্রুপ কয়টা সাব-কলামে ভাগ হবে তা লিখুন (কমা দিয়ে আলাদা)। "
    "ফাঁকা রাখলে সব গ্রুপে ডিফল্ট ২টা সাব-কলাম হবে। যেমন: 2,2,3"
)
group_subcols_text = st.text_input("সাব-কলাম সংখ্যা (অপশনাল)", value="")

st.caption(
    "🔧 অপশনাল: কোনো রোল নাম্বার অনুপস্থিত/বাদ দিতে চাইলে এখানে কমা দিয়ে লিখুন। "
    "একক রোল বা হাইফেন দিয়ে রেঞ্জ দুটোই দেওয়া যাবে। ওই রোলগুলোর জায়গায় কোনো বক্স "
    "থাকবে না, পরের রোল উপরে উঠে আসবে। যেমন: 102175-102180,102185,102190-102200"
)
absent_rolls_text = st.text_input("অনুপস্থিত রোল নাম্বার (অপশনাল)", value="")

# ---- ইনপুট পার্স ও ভ্যালিডেশন ----
parse_ok = True
group_counts, group_subcols, absent_rolls = [], [], set()

try:
    group_counts = [int(x.strip()) for x in group_counts_text.split(",") if x.strip()]
    if not group_counts:
        raise ValueError
except ValueError:
    parse_ok = False
    st.error("কলাম বিন্যাস সঠিকভাবে লিখুন, শুধু সংখ্যা ও কমা ব্যবহার করুন। যেমন: 12,12,16")

if parse_ok:
    if group_subcols_text.strip():
        try:
            group_subcols = [int(x.strip()) for x in group_subcols_text.split(",") if x.strip()]
            if len(group_subcols) != len(group_counts):
                parse_ok = False
                st.error(
                    f"সাব-কলাম সংখ্যার লিস্টে {len(group_subcols)}টা মান আছে, "
                    f"কিন্তু কলাম-গ্রুপ আছে {len(group_counts)}টা। দুটোর সংখ্যা মিলতে হবে, "
                    f"অথবা পুরো ফিল্ডটা ফাঁকা রাখুন ডিফল্ট ২ ব্যবহারের জন্য।"
                )
            elif any(s < 1 for s in group_subcols):
                parse_ok = False
                st.error("সাব-কলাম সংখ্যা কমপক্ষে ১ হতে হবে।")
        except ValueError:
            parse_ok = False
            st.error("সাব-কলাম সংখ্যা সঠিকভাবে লিখুন, শুধু সংখ্যা ও কমা। যেমন: 2,2,3")
    else:
        group_subcols = [2] * len(group_counts)  # ডিফল্ট: প্রতি গ্রুপে ২টা সাব-কলাম

if absent_rolls_text.strip():
    try:
        absent_rolls = parse_absent_rolls(absent_rolls_text)
    except ValueError:
        parse_ok = False
        st.error(
            "অনুপস্থিত রোল নাম্বার সঠিকভাবে লিখুন। একক রোল বা হাইফেন-রেঞ্জ, কমা দিয়ে "
            "আলাদা করুন। যেমন: 102175-102180,102185,102190-102200"
        )

if parse_ok:
    total_from_groups = sum(group_counts)
    if total_from_groups != total_students:
        st.warning(
            f"⚠️ আপনার কলাম বিন্যাস অনুযায়ী মোট হচ্ছে {total_from_groups} জন, "
            f"কিন্তু মোট পরীক্ষার্থী লিখেছেন {total_students} জন। "
            f"মিলিয়ে নিন (নাহলে ফাইলে {total_from_groups} জনের রোলই বসবে)।"
        )
    if absent_rolls:
        st.info(f"ℹ️ মোট {len(absent_rolls)}টা রোল নাম্বার অনুপস্থিত হিসেবে বাদ দেওয়া হবে।")

if st.button("📄 Word ফাইল তৈরি করুন", type="primary", disabled=not parse_ok):
    docx_bytes, last_roll = build_document(
        room_no=room_no,
        start_roll=int(start_roll),
        total_students=int(total_students),
        group_counts=group_counts,
        group_subcols=group_subcols,
        absent_rolls=absent_rolls,
        gap_width_cm=float(gap_width),
    )
    st.success(f"তৈরি হয়েছে! রোল নাম্বার রেঞ্জ: {int(start_roll)} থেকে {last_roll} পর্যন্ত")
    st.download_button(
        label="⬇️ ডাউনলোড করুন (seating_sheet.docx)",
        data=docx_bytes,
        file_name=f"seating_sheet_room_{room_no}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )